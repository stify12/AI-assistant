"""
日报服务模块 (US-14)

提供测试日报的自动生成功能，包括：
- 每日测试日报生成
- 新增错误类型检测
- 高频错误题目统计
- 日报导出 (Word格式)
- 历史日报查询

遵循 NFR-34 代码质量标准
"""
import os
import uuid
import json
from datetime import datetime, timedelta
from typing import Optional, Dict, Any, List, Set

from .database_service import AppDatabaseService
from .storage_service import StorageService
from .llm_service import LLMService


# 学科ID映射
SUBJECT_MAP = {
    0: '英语',
    1: '语文',
    2: '数学',
    3: '物理',
    4: '化学',
    5: '生物',
    6: '地理'
}

# 默认AI模型版本
DEFAULT_MODEL_VERSION = 'doubao-1-5-vision-pro-32k-250115'


class ReportService:
    """
    日报服务类 (US-14)
    
    提供测试日报的生成、查询和导出功能。
    支持自动生成和手动触发两种方式。
    
    日报内容包括：
    - 今日任务完成数
    - 准确率及变化（与昨日、上周同日对比）
    - 主要错误类型 Top 5
    - 新增错误类型
    - 高频错误题目
    - 明日计划
    - 异常情况
    - AI总结
    """
    
    # 导出目录
    EXPORTS_DIR = 'exports'
    
    @staticmethod
    def generate_daily_report(
        date: str = None,
        generated_by: str = 'manual'
    ) -> Dict[str, Any]:
        """
        生成测试日报 (9.1.1)
        
        生成指定日期的测试日报，包含任务完成情况、准确率变化、
        错误分析、AI总结等内容。
        
        Args:
            date: 日报日期，格式 YYYY-MM-DD，默认今天
            generated_by: 生成方式 auto|manual，默认 manual
            
        Returns:
            dict: 日报数据，包含以下字段：
                - report_id: 日报唯一标识
                - report_date: 日报日期
                - task_completed: 完成任务数
                - task_planned: 计划任务数
                - accuracy: 当日准确率
                - accuracy_change: 与昨日对比变化
                - accuracy_week_change: 与上周同日对比变化
                - top_errors: 主要错误类型 Top 5
                - new_error_types: 新增错误类型
                - high_freq_errors: 高频错误题目
                - tomorrow_plan: 明日计划
                - anomalies: 异常情况
                - model_version: AI模型版本
                - ai_summary: AI生成的总结
                - raw_content: 完整日报内容（Markdown）
                
        Raises:
            ValueError: 日期格式错误
        """
        # 解析日期参数
        if date:
            try:
                report_date = datetime.strptime(date, '%Y-%m-%d').date()
            except ValueError:
                raise ValueError(f'无效的日期格式: {date}，应为 YYYY-MM-DD')
        else:
            report_date = datetime.now().date()
        
        report_date_str = report_date.strftime('%Y-%m-%d')
        
        # 检查是否已存在该日期的日报
        existing_report = ReportService._get_report_by_date(report_date_str)
        if existing_report:
            # 如果已存在，返回现有日报
            return existing_report
        
        # 生成日报ID
        report_id = str(uuid.uuid4())[:8]
        
        # 1. 获取当日任务统计
        task_stats = ReportService._get_task_stats(report_date)
        
        # 2. 计算准确率及变化
        accuracy_stats = ReportService._calculate_accuracy_stats(report_date)
        
        # 3. 获取主要错误类型 Top 5
        top_errors = ReportService._get_top_errors(report_date)
        
        # 4. 获取新增错误类型
        new_error_types = ReportService.get_new_error_types(report_date_str)
        
        # 5. 获取高频错误题目
        high_freq_errors = ReportService.get_high_freq_errors(report_date_str)
        
        # 6. 获取明日计划（待处理任务）
        tomorrow_plan = ReportService._get_tomorrow_plan()
        
        # 7. 检测异常情况
        anomalies = ReportService._detect_anomalies(report_date, accuracy_stats)
        
        # 8. 获取模型版本信息
        model_version = ReportService._get_model_version(report_date)
        
        # 9. 生成AI总结
        ai_summary = ReportService._generate_ai_summary(
            report_date_str,
            task_stats,
            accuracy_stats,
            top_errors,
            new_error_types,
            high_freq_errors,
            anomalies
        )
        
        # 10. 生成完整日报内容（Markdown）
        raw_content = ReportService._generate_raw_content(
            report_date_str,
            task_stats,
            accuracy_stats,
            top_errors,
            new_error_types,
            high_freq_errors,
            tomorrow_plan,
            anomalies,
            model_version,
            ai_summary
        )
        
        # 构建日报数据
        report_data = {
            'report_id': report_id,
            'report_date': report_date_str,
            'task_completed': task_stats.get('completed', 0),
            'task_planned': task_stats.get('planned', 0),
            'accuracy': accuracy_stats.get('current', 0),
            'accuracy_change': accuracy_stats.get('day_change', 0),
            'accuracy_week_change': accuracy_stats.get('week_change', 0),
            'top_errors': top_errors,
            'new_error_types': new_error_types,
            'high_freq_errors': high_freq_errors,
            'tomorrow_plan': tomorrow_plan,
            'anomalies': anomalies,
            'model_version': model_version,
            'ai_summary': ai_summary,
            'raw_content': raw_content,
            'generated_by': generated_by
        }
        
        # 保存到数据库
        ReportService._save_report(report_data)
        
        return report_data
    
    @staticmethod
    def get_new_error_types(date: str) -> List[str]:
        """
        获取新增错误类型 (9.1.2)
        
        对比今日错误类型与历史错误类型，找出今日首次出现的错误模式。
        
        Args:
            date: 日期，格式 YYYY-MM-DD
            
        Returns:
            list: 新增错误类型列表
        """
        try:
            report_date = datetime.strptime(date, '%Y-%m-%d').date()
        except ValueError:
            return []
        
        # 获取今日错误类型
        today_errors = ReportService._get_error_types_for_date(report_date)
        
        # 获取历史错误类型（过去30天）
        history_start = report_date - timedelta(days=30)
        history_errors = ReportService._get_error_types_in_range(history_start, report_date - timedelta(days=1))
        
        # 找出新增的错误类型
        new_types = [et for et in today_errors if et not in history_errors]
        
        return new_types
    
    @staticmethod
    def get_high_freq_errors(date: str, min_count: int = 3) -> List[Dict[str, Any]]:
        """
        获取高频错误题目 (9.1.3)
        
        统计同一题目多次出错的情况，返回出错次数>=min_count的题目。
        
        Args:
            date: 日期，格式 YYYY-MM-DD
            min_count: 最小出错次数阈值，默认3
            
        Returns:
            list: 高频错误题目列表，每个元素包含：
                - index: 题号
                - count: 出错次数
                - book_name: 书本名称
                - page_num: 页码
                - error_types: 错误类型列表
        """
        try:
            report_date = datetime.strptime(date, '%Y-%m-%d').date()
        except ValueError:
            return []
        
        # 加载当日任务
        tasks = ReportService._load_tasks_for_date(report_date)
        
        # 统计每个题目的错误次数
        # key: (book_id, page_num, index)
        error_counts: Dict[tuple, Dict[str, Any]] = {}
        
        for task in tasks:
            for hw_item in task.get('homework_items', []):
                evaluation = hw_item.get('evaluation') or {}
                errors = evaluation.get('errors') or []
                book_id = hw_item.get('book_id', '')
                book_name = hw_item.get('book_name', '')
                page_num = hw_item.get('page_num', 0)
                
                for error in errors:
                    index = error.get('index', '')
                    error_type = error.get('error_type', '')
                    
                    key = (book_id, page_num, index)
                    if key not in error_counts:
                        error_counts[key] = {
                            'index': index,
                            'count': 0,
                            'book_id': book_id,
                            'book_name': book_name,
                            'page_num': page_num,
                            'error_types': set()
                        }
                    
                    error_counts[key]['count'] += 1
                    error_counts[key]['error_types'].add(error_type)
        
        # 筛选高频错误（出错次数>=min_count）
        high_freq = []
        for key, data in error_counts.items():
            if data['count'] >= min_count:
                high_freq.append({
                    'index': data['index'],
                    'count': data['count'],
                    'book_name': data['book_name'],
                    'page_num': data['page_num'],
                    'error_types': list(data['error_types'])
                })
        
        # 按出错次数降序排列
        high_freq.sort(key=lambda x: x['count'], reverse=True)
        
        # 返回前10个
        return high_freq[:10]
    
    @staticmethod
    def export_report(report_id: str, format: str = 'docx') -> str:
        """
        导出日报 (9.1.4)
        
        将日报导出为指定格式的文件。
        
        Args:
            report_id: 日报ID
            format: 导出格式 docx，默认 docx
            
        Returns:
            str: 导出文件的路径
            
        Raises:
            ValueError: 日报不存在或格式不支持
        """
        # 获取日报数据
        report = ReportService._get_report_by_id(report_id)
        if not report:
            raise ValueError(f'日报不存在: {report_id}')
        
        # 确保导出目录存在
        StorageService.ensure_dir(ReportService.EXPORTS_DIR)
        
        if format == 'docx':
            return ReportService._export_to_docx(report)
        else:
            raise ValueError(f'不支持的导出格式: {format}，目前仅支持 docx')
    
    @staticmethod
    def get_report_history(
        page: int = 1,
        page_size: int = 30
    ) -> Dict[str, Any]:
        """
        获取历史日报列表 (9.1.5)
        
        分页查询历史日报，按日期倒序排列。
        
        Args:
            page: 页码，从1开始
            page_size: 每页数量，默认30
            
        Returns:
            dict: 包含以下字段：
                - reports: 日报列表
                - pagination: 分页信息
        """
        result = {
            'reports': [],
            'pagination': {
                'page': page,
                'page_size': page_size,
                'total': 0,
                'total_pages': 0
            }
        }
        
        try:
            # 查询总数
            count_sql = "SELECT COUNT(*) as total FROM daily_reports"
            count_result = AppDatabaseService.execute_query(count_sql)
            total = count_result[0]['total'] if count_result else 0
            
            result['pagination']['total'] = total
            result['pagination']['total_pages'] = (total + page_size - 1) // page_size if page_size > 0 else 0
            
            # 查询日报列表
            offset = (page - 1) * page_size
            sql = """
                SELECT report_id, report_date, task_completed, task_planned,
                       accuracy, accuracy_change, accuracy_week_change,
                       generated_by, created_at
                FROM daily_reports
                ORDER BY report_date DESC
                LIMIT %s OFFSET %s
            """
            reports = AppDatabaseService.execute_query(sql, (page_size, offset))
            
            if reports:
                for report in reports:
                    # 格式化日期
                    report_date = report.get('report_date')
                    if report_date:
                        report['report_date'] = report_date.strftime('%Y-%m-%d') if hasattr(report_date, 'strftime') else str(report_date)
                    
                    created_at = report.get('created_at')
                    if created_at:
                        report['created_at'] = created_at.isoformat() if hasattr(created_at, 'isoformat') else str(created_at)
                    
                    result['reports'].append(report)
            
        except Exception as e:
            print(f"[ReportService] 获取历史日报列表失败: {e}")
        
        return result

    
    # ========== 私有辅助方法 ==========
    
    @staticmethod
    def _get_report_by_date(date_str: str) -> Optional[Dict[str, Any]]:
        """
        根据日期获取日报
        
        Args:
            date_str: 日期字符串 YYYY-MM-DD
            
        Returns:
            dict: 日报数据，不存在返回 None
        """
        try:
            sql = """
                SELECT * FROM daily_reports WHERE report_date = %s
            """
            results = AppDatabaseService.execute_query(sql, (date_str,))
            
            if results:
                report = results[0]
                return ReportService._format_report(report)
            
        except Exception as e:
            print(f"[ReportService] 查询日报失败: {e}")
        
        return None
    
    @staticmethod
    def _get_report_by_id(report_id: str) -> Optional[Dict[str, Any]]:
        """
        根据ID获取日报
        
        Args:
            report_id: 日报ID
            
        Returns:
            dict: 日报数据，不存在返回 None
        """
        try:
            sql = """
                SELECT * FROM daily_reports WHERE report_id = %s
            """
            results = AppDatabaseService.execute_query(sql, (report_id,))
            
            if results:
                report = results[0]
                return ReportService._format_report(report)
            
        except Exception as e:
            print(f"[ReportService] 查询日报失败: {e}")
        
        return None
    
    @staticmethod
    def _format_report(report: Dict[str, Any]) -> Dict[str, Any]:
        """
        格式化日报数据
        
        将数据库查询结果转换为API响应格式。
        
        Args:
            report: 数据库查询结果
            
        Returns:
            dict: 格式化后的日报数据
        """
        # 解析JSON字段
        json_fields = ['top_errors', 'new_error_types', 'high_freq_errors', 
                       'tomorrow_plan', 'anomalies']
        
        for field in json_fields:
            value = report.get(field)
            if value and isinstance(value, str):
                try:
                    report[field] = json.loads(value)
                except:
                    report[field] = []
            elif value is None:
                report[field] = []
        
        # 格式化日期
        report_date = report.get('report_date')
        if report_date:
            report['report_date'] = report_date.strftime('%Y-%m-%d') if hasattr(report_date, 'strftime') else str(report_date)
        
        created_at = report.get('created_at')
        if created_at:
            report['created_at'] = created_at.isoformat() if hasattr(created_at, 'isoformat') else str(created_at)
        
        return report
    
    @staticmethod
    def _save_report(report_data: Dict[str, Any]) -> None:
        """
        保存日报到数据库
        
        Args:
            report_data: 日报数据
        """
        try:
            sql = """
                INSERT INTO daily_reports 
                (report_id, report_date, task_completed, task_planned,
                 accuracy, accuracy_change, accuracy_week_change,
                 top_errors, new_error_types, high_freq_errors,
                 tomorrow_plan, anomalies, model_version,
                 ai_summary, raw_content, generated_by, created_at)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                ON DUPLICATE KEY UPDATE
                task_completed = VALUES(task_completed),
                task_planned = VALUES(task_planned),
                accuracy = VALUES(accuracy),
                accuracy_change = VALUES(accuracy_change),
                accuracy_week_change = VALUES(accuracy_week_change),
                top_errors = VALUES(top_errors),
                new_error_types = VALUES(new_error_types),
                high_freq_errors = VALUES(high_freq_errors),
                tomorrow_plan = VALUES(tomorrow_plan),
                anomalies = VALUES(anomalies),
                model_version = VALUES(model_version),
                ai_summary = VALUES(ai_summary),
                raw_content = VALUES(raw_content),
                generated_by = VALUES(generated_by)
            """
            
            AppDatabaseService.execute_insert(sql, (
                report_data['report_id'],
                report_data['report_date'],
                report_data['task_completed'],
                report_data['task_planned'],
                report_data['accuracy'],
                report_data['accuracy_change'],
                report_data['accuracy_week_change'],
                json.dumps(report_data['top_errors'], ensure_ascii=False),
                json.dumps(report_data['new_error_types'], ensure_ascii=False),
                json.dumps(report_data['high_freq_errors'], ensure_ascii=False),
                json.dumps(report_data['tomorrow_plan'], ensure_ascii=False),
                json.dumps(report_data['anomalies'], ensure_ascii=False),
                report_data['model_version'],
                report_data['ai_summary'],
                report_data['raw_content'],
                report_data['generated_by'],
                datetime.now()
            ))
            
        except Exception as e:
            print(f"[ReportService] 保存日报失败: {e}")
            raise
    
    @staticmethod
    def _load_tasks_for_date(date: datetime.date) -> List[Dict[str, Any]]:
        """
        加载指定日期的批量任务
        
        Args:
            date: 日期
            
        Returns:
            list: 任务列表
        """
        tasks = []
        batch_tasks_dir = StorageService.BATCH_TASKS_DIR
        
        # 日期范围：当天 00:00 到 23:59:59
        date_start = datetime.combine(date, datetime.min.time())
        date_end = datetime.combine(date, datetime.max.time())
        
        try:
            StorageService.ensure_dir(batch_tasks_dir)
            
            for filename in os.listdir(batch_tasks_dir):
                if filename.endswith('.json'):
                    filepath = os.path.join(batch_tasks_dir, filename)
                    try:
                        with open(filepath, 'r', encoding='utf-8') as f:
                            task_data = json.load(f)
                            
                            # 检查创建时间是否在指定日期
                            created_at = task_data.get('created_at', '')
                            if created_at:
                                try:
                                    task_time = datetime.fromisoformat(created_at.replace('Z', '+00:00'))
                                    if task_time.tzinfo:
                                        task_time = task_time.replace(tzinfo=None)
                                    
                                    if date_start <= task_time <= date_end:
                                        tasks.append(task_data)
                                except:
                                    continue
                    except Exception as e:
                        print(f"[ReportService] 加载任务文件失败 {filename}: {e}")
                        continue
            
        except Exception as e:
            print(f"[ReportService] 扫描批量任务目录失败: {e}")
        
        return tasks
    
    @staticmethod
    def _get_task_stats(date: datetime.date) -> Dict[str, int]:
        """
        获取任务统计
        
        Args:
            date: 日期
            
        Returns:
            dict: {completed, planned, total_questions, correct_questions}
        """
        tasks = ReportService._load_tasks_for_date(date)
        
        completed = 0
        planned = 0
        total_questions = 0
        correct_questions = 0
        
        for task in tasks:
            status = task.get('status', 'pending')
            if status == 'completed':
                completed += 1
            planned += 1
            
            # 统计题目数
            overall_report = task.get('overall_report') or {}
            total_questions += overall_report.get('total_questions', 0)
            correct_questions += overall_report.get('correct_questions', 0)
        
        return {
            'completed': completed,
            'planned': planned,
            'total_questions': total_questions,
            'correct_questions': correct_questions
        }
    
    @staticmethod
    def _calculate_accuracy_stats(date: datetime.date) -> Dict[str, float]:
        """
        计算准确率统计
        
        Args:
            date: 日期
            
        Returns:
            dict: {current, yesterday, last_week, day_change, week_change}
        """
        # 当日准确率
        today_stats = ReportService._get_task_stats(date)
        current = 0
        if today_stats['total_questions'] > 0:
            current = round(today_stats['correct_questions'] / today_stats['total_questions'], 4)
        
        # 昨日准确率
        yesterday = date - timedelta(days=1)
        yesterday_stats = ReportService._get_task_stats(yesterday)
        yesterday_acc = 0
        if yesterday_stats['total_questions'] > 0:
            yesterday_acc = round(yesterday_stats['correct_questions'] / yesterday_stats['total_questions'], 4)
        
        # 上周同日准确率
        last_week = date - timedelta(days=7)
        last_week_stats = ReportService._get_task_stats(last_week)
        last_week_acc = 0
        if last_week_stats['total_questions'] > 0:
            last_week_acc = round(last_week_stats['correct_questions'] / last_week_stats['total_questions'], 4)
        
        return {
            'current': current,
            'yesterday': yesterday_acc,
            'last_week': last_week_acc,
            'day_change': round(current - yesterday_acc, 4),
            'week_change': round(current - last_week_acc, 4)
        }
    
    @staticmethod
    def _get_top_errors(date: datetime.date, limit: int = 5) -> List[Dict[str, Any]]:
        """
        获取主要错误类型 Top N
        
        Args:
            date: 日期
            limit: 返回数量
            
        Returns:
            list: [{type, count}]
        """
        tasks = ReportService._load_tasks_for_date(date)
        
        # 统计错误类型
        error_counts: Dict[str, int] = {}
        
        for task in tasks:
            for hw_item in task.get('homework_items', []):
                evaluation = hw_item.get('evaluation') or {}
                errors = evaluation.get('errors') or []
                
                for error in errors:
                    error_type = error.get('error_type', '其他')
                    error_counts[error_type] = error_counts.get(error_type, 0) + 1
        
        # 排序并返回 Top N
        sorted_errors = sorted(error_counts.items(), key=lambda x: x[1], reverse=True)
        
        return [{'type': et, 'count': count} for et, count in sorted_errors[:limit]]

    
    @staticmethod
    def _get_error_types_for_date(date: datetime.date) -> Set[str]:
        """
        获取指定日期的所有错误类型
        
        Args:
            date: 日期
            
        Returns:
            set: 错误类型集合
        """
        tasks = ReportService._load_tasks_for_date(date)
        error_types = set()
        
        for task in tasks:
            for hw_item in task.get('homework_items', []):
                evaluation = hw_item.get('evaluation') or {}
                errors = evaluation.get('errors') or []
                
                for error in errors:
                    error_type = error.get('error_type', '')
                    if error_type:
                        error_types.add(error_type)
        
        return error_types
    
    @staticmethod
    def _get_error_types_in_range(start_date: datetime.date, end_date: datetime.date) -> Set[str]:
        """
        获取日期范围内的所有错误类型
        
        Args:
            start_date: 开始日期
            end_date: 结束日期
            
        Returns:
            set: 错误类型集合
        """
        error_types = set()
        current_date = start_date
        
        while current_date <= end_date:
            day_errors = ReportService._get_error_types_for_date(current_date)
            error_types.update(day_errors)
            current_date += timedelta(days=1)
        
        return error_types
    
    @staticmethod
    def _get_tomorrow_plan() -> List[Dict[str, Any]]:
        """
        获取明日计划（待处理任务）
        
        Returns:
            list: 待处理任务列表
        """
        plans = []
        
        try:
            # 查询状态为 active 的测试计划
            sql = """
                SELECT plan_id, name, target_count, completed_count
                FROM test_plans
                WHERE status = 'active'
                ORDER BY created_at DESC
                LIMIT 5
            """
            results = AppDatabaseService.execute_query(sql)
            
            if results:
                for plan in results:
                    plans.append({
                        'plan_id': plan.get('plan_id'),
                        'name': plan.get('name'),
                        'target_count': plan.get('target_count', 0),
                        'completed_count': plan.get('completed_count', 0),
                        'remaining': plan.get('target_count', 0) - plan.get('completed_count', 0)
                    })
            
        except Exception as e:
            print(f"[ReportService] 获取明日计划失败: {e}")
        
        return plans
    
    @staticmethod
    def _detect_anomalies(date: datetime.date, accuracy_stats: Dict[str, float]) -> List[Dict[str, Any]]:
        """
        检测异常情况
        
        检测准确率异常波动、任务失败等情况。
        
        Args:
            date: 日期
            accuracy_stats: 准确率统计
            
        Returns:
            list: 异常情况列表
        """
        anomalies = []
        
        # 1. 检测准确率大幅下降（下降超过10%）
        day_change = accuracy_stats.get('day_change', 0)
        if day_change < -0.1:
            anomalies.append({
                'type': 'accuracy_drop',
                'severity': 'high',
                'message': f'准确率较昨日下降 {abs(day_change)*100:.1f}%',
                'value': day_change
            })
        
        # 2. 检测失败任务
        tasks = ReportService._load_tasks_for_date(date)
        failed_tasks = [t for t in tasks if t.get('status') == 'failed']
        if failed_tasks:
            anomalies.append({
                'type': 'task_failed',
                'severity': 'medium',
                'message': f'有 {len(failed_tasks)} 个任务执行失败',
                'value': len(failed_tasks)
            })
        
        # 3. 检测准确率过低（低于60%）
        current_acc = accuracy_stats.get('current', 0)
        if current_acc > 0 and current_acc < 0.6:
            anomalies.append({
                'type': 'low_accuracy',
                'severity': 'high',
                'message': f'当日准确率仅 {current_acc*100:.1f}%，低于60%警戒线',
                'value': current_acc
            })
        
        return anomalies
    
    @staticmethod
    def _get_model_version(date: datetime.date) -> str:
        """
        获取当日使用的AI模型版本
        
        Args:
            date: 日期
            
        Returns:
            str: 模型版本
        """
        # 从配置或任务中获取模型版本
        # 这里返回默认值，实际可以从任务数据中提取
        return DEFAULT_MODEL_VERSION
    
    @staticmethod
    def _generate_ai_summary(
        date_str: str,
        task_stats: Dict[str, int],
        accuracy_stats: Dict[str, float],
        top_errors: List[Dict[str, Any]],
        new_error_types: List[str],
        high_freq_errors: List[Dict[str, Any]],
        anomalies: List[Dict[str, Any]]
    ) -> str:
        """
        使用AI生成日报总结
        
        调用 DeepSeek V3.2 模型生成日报的关键信息总结。
        
        Args:
            date_str: 日期字符串
            task_stats: 任务统计
            accuracy_stats: 准确率统计
            top_errors: 主要错误类型
            new_error_types: 新增错误类型
            high_freq_errors: 高频错误题目
            anomalies: 异常情况
            
        Returns:
            str: AI生成的总结文本
        """
        # 构建提示词
        prompt = f"""请根据以下测试数据，生成一段简洁的测试日报总结（100-200字）：

日期：{date_str}

任务完成情况：
- 完成任务数：{task_stats.get('completed', 0)}
- 计划任务数：{task_stats.get('planned', 0)}
- 测试题目数：{task_stats.get('total_questions', 0)}
- 正确题目数：{task_stats.get('correct_questions', 0)}

准确率情况：
- 当日准确率：{accuracy_stats.get('current', 0)*100:.1f}%
- 较昨日变化：{accuracy_stats.get('day_change', 0)*100:+.1f}%
- 较上周同日变化：{accuracy_stats.get('week_change', 0)*100:+.1f}%

主要错误类型 Top 5：
{json.dumps(top_errors, ensure_ascii=False, indent=2) if top_errors else '无'}

新增错误类型：
{', '.join(new_error_types) if new_error_types else '无'}

高频错误题目（出错>=3次）：
{json.dumps(high_freq_errors[:3], ensure_ascii=False, indent=2) if high_freq_errors else '无'}

异常情况：
{json.dumps(anomalies, ensure_ascii=False, indent=2) if anomalies else '无'}

请用专业、客观的语气总结今日测试情况，包括：
1. 整体表现评价
2. 主要问题分析
3. 改进建议（如有）

注意：不要使用emoji，保持简洁专业。"""

        system_prompt = '你是一个专业的AI批改效果测试分析师，擅长分析测试数据并给出专业的总结报告。请用中文回答。'
        
        try:
            # 调用 DeepSeek 生成总结
            result = LLMService.call_deepseek(
                prompt=prompt,
                system_prompt=system_prompt,
                model='deepseek-v3.2',
                timeout=30
            )
            
            if result.get('success') and result.get('content'):
                return result['content'].strip()
            else:
                error_msg = result.get('error', '未知错误')
                print(f"[ReportService] AI生成总结失败: {error_msg}")
                return f"今日完成 {task_stats.get('completed', 0)} 个测试任务，准确率 {accuracy_stats.get('current', 0)*100:.1f}%。"
                
        except Exception as e:
            print(f"[ReportService] AI生成总结异常: {e}")
            return f"今日完成 {task_stats.get('completed', 0)} 个测试任务，准确率 {accuracy_stats.get('current', 0)*100:.1f}%。"
    
    @staticmethod
    def _generate_raw_content(
        date_str: str,
        task_stats: Dict[str, int],
        accuracy_stats: Dict[str, float],
        top_errors: List[Dict[str, Any]],
        new_error_types: List[str],
        high_freq_errors: List[Dict[str, Any]],
        tomorrow_plan: List[Dict[str, Any]],
        anomalies: List[Dict[str, Any]],
        model_version: str,
        ai_summary: str
    ) -> str:
        """
        生成完整日报内容（Markdown格式）
        
        Args:
            各项统计数据
            
        Returns:
            str: Markdown格式的完整日报
        """
        content = f"""# 测试日报 {date_str}

## AI总结

{ai_summary}

---

## 一、任务完成情况

| 指标 | 数值 |
|------|------|
| 完成任务数 | {task_stats.get('completed', 0)} |
| 计划任务数 | {task_stats.get('planned', 0)} |
| 测试题目数 | {task_stats.get('total_questions', 0)} |
| 正确题目数 | {task_stats.get('correct_questions', 0)} |

## 二、准确率分析

| 指标 | 数值 |
|------|------|
| 当日准确率 | {accuracy_stats.get('current', 0)*100:.2f}% |
| 较昨日变化 | {accuracy_stats.get('day_change', 0)*100:+.2f}% |
| 较上周同日变化 | {accuracy_stats.get('week_change', 0)*100:+.2f}% |

## 三、主要错误类型 Top 5

| 排名 | 错误类型 | 出现次数 |
|------|----------|----------|
"""
        
        # 添加错误类型表格
        for i, error in enumerate(top_errors, 1):
            content += f"| {i} | {error.get('type', '')} | {error.get('count', 0)} |\n"
        
        if not top_errors:
            content += "| - | 无错误记录 | - |\n"
        
        # 新增错误类型
        content += "\n## 四、新增错误类型\n\n"
        if new_error_types:
            for et in new_error_types:
                content += f"- **{et}** (今日首次出现)\n"
        else:
            content += "无新增错误类型\n"
        
        # 高频错误题目
        content += "\n## 五、高频错误题目\n\n"
        if high_freq_errors:
            content += "| 题号 | 书本 | 页码 | 出错次数 | 错误类型 |\n"
            content += "|------|------|------|----------|----------|\n"
            for error in high_freq_errors[:5]:
                error_types_str = ', '.join(error.get('error_types', []))
                content += f"| {error.get('index', '')} | {error.get('book_name', '')} | {error.get('page_num', '')} | {error.get('count', 0)} | {error_types_str} |\n"
        else:
            content += "无高频错误题目（出错次数>=3）\n"
        
        # 异常情况
        content += "\n## 六、异常情况\n\n"
        if anomalies:
            for anomaly in anomalies:
                severity_icon = "!" if anomaly.get('severity') == 'high' else "-"
                content += f"- [{severity_icon}] {anomaly.get('message', '')}\n"
        else:
            content += "无异常情况\n"
        
        # 明日计划
        content += "\n## 七、明日计划\n\n"
        if tomorrow_plan:
            for plan in tomorrow_plan:
                remaining = plan.get('remaining', 0)
                content += f"- {plan.get('name', '')}: 剩余 {remaining} 个任务\n"
        else:
            content += "暂无待处理计划\n"
        
        # 模型版本
        content += f"\n---\n\n**模型版本**: {model_version}\n"
        content += f"**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        
        return content

    
    @staticmethod
    def _export_to_docx(report: Dict[str, Any]) -> str:
        """
        导出日报为Word文档
        
        使用 python-docx 生成 Word 文档。
        
        Args:
            report: 日报数据
            
        Returns:
            str: 导出文件路径
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ValueError('python-docx 未安装，无法导出 Word 文档')
        
        # 创建文档
        doc = Document()
        
        # 设置标题
        report_date = report.get('report_date', '')
        title = doc.add_heading(f'测试日报 {report_date}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # AI总结
        doc.add_heading('AI总结', level=1)
        ai_summary = report.get('ai_summary', '')
        doc.add_paragraph(ai_summary)
        
        # 任务完成情况
        doc.add_heading('一、任务完成情况', level=1)
        task_table = doc.add_table(rows=5, cols=2)
        task_table.style = 'Table Grid'
        
        task_data = [
            ('指标', '数值'),
            ('完成任务数', str(report.get('task_completed', 0))),
            ('计划任务数', str(report.get('task_planned', 0))),
            ('当日准确率', f"{report.get('accuracy', 0)*100:.2f}%"),
            ('较昨日变化', f"{report.get('accuracy_change', 0)*100:+.2f}%")
        ]
        
        for i, (label, value) in enumerate(task_data):
            task_table.rows[i].cells[0].text = label
            task_table.rows[i].cells[1].text = value
        
        # 主要错误类型
        doc.add_heading('二、主要错误类型 Top 5', level=1)
        top_errors = report.get('top_errors', [])
        if top_errors:
            error_table = doc.add_table(rows=len(top_errors)+1, cols=3)
            error_table.style = 'Table Grid'
            
            # 表头
            error_table.rows[0].cells[0].text = '排名'
            error_table.rows[0].cells[1].text = '错误类型'
            error_table.rows[0].cells[2].text = '出现次数'
            
            for i, error in enumerate(top_errors, 1):
                error_table.rows[i].cells[0].text = str(i)
                error_table.rows[i].cells[1].text = error.get('type', '')
                error_table.rows[i].cells[2].text = str(error.get('count', 0))
        else:
            doc.add_paragraph('无错误记录')
        
        # 新增错误类型
        doc.add_heading('三、新增错误类型', level=1)
        new_error_types = report.get('new_error_types', [])
        if new_error_types:
            for et in new_error_types:
                doc.add_paragraph(f'• {et} (今日首次出现)', style='List Bullet')
        else:
            doc.add_paragraph('无新增错误类型')
        
        # 高频错误题目
        doc.add_heading('四、高频错误题目', level=1)
        high_freq_errors = report.get('high_freq_errors', [])
        if high_freq_errors:
            hf_table = doc.add_table(rows=len(high_freq_errors)+1, cols=4)
            hf_table.style = 'Table Grid'
            
            # 表头
            hf_table.rows[0].cells[0].text = '题号'
            hf_table.rows[0].cells[1].text = '书本'
            hf_table.rows[0].cells[2].text = '页码'
            hf_table.rows[0].cells[3].text = '出错次数'
            
            for i, error in enumerate(high_freq_errors[:5], 1):
                hf_table.rows[i].cells[0].text = str(error.get('index', ''))
                hf_table.rows[i].cells[1].text = error.get('book_name', '')
                hf_table.rows[i].cells[2].text = str(error.get('page_num', ''))
                hf_table.rows[i].cells[3].text = str(error.get('count', 0))
        else:
            doc.add_paragraph('无高频错误题目')
        
        # 异常情况
        doc.add_heading('五、异常情况', level=1)
        anomalies = report.get('anomalies', [])
        if anomalies:
            for anomaly in anomalies:
                severity = '[!]' if anomaly.get('severity') == 'high' else '[-]'
                doc.add_paragraph(f'{severity} {anomaly.get("message", "")}')
        else:
            doc.add_paragraph('无异常情况')
        
        # 明日计划
        doc.add_heading('六、明日计划', level=1)
        tomorrow_plan = report.get('tomorrow_plan', [])
        if tomorrow_plan:
            for plan in tomorrow_plan:
                remaining = plan.get('remaining', 0)
                doc.add_paragraph(f'• {plan.get("name", "")}: 剩余 {remaining} 个任务', style='List Bullet')
        else:
            doc.add_paragraph('暂无待处理计划')
        
        # 页脚信息
        doc.add_paragraph('')
        footer = doc.add_paragraph()
        footer.add_run(f'模型版本: {report.get("model_version", "")}').italic = True
        footer.add_run(f'  |  生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}').italic = True
        
        # 保存文件
        filename = f'daily_report_{report_date}_{report.get("report_id", "")}.docx'
        filepath = os.path.join(ReportService.EXPORTS_DIR, filename)
        doc.save(filepath)
        
        return filepath
    
    @staticmethod
    def delete_old_reports(days: int = 30) -> int:
        """
        删除过期日报
        
        删除超过指定天数的历史日报。
        
        Args:
            days: 保留天数，默认30天
            
        Returns:
            int: 删除的日报数量
        """
        try:
            cutoff_date = (datetime.now() - timedelta(days=days)).strftime('%Y-%m-%d')
            
            sql = """
                DELETE FROM daily_reports WHERE report_date < %s
            """
            result = AppDatabaseService.execute_update(sql, (cutoff_date,))
            
            deleted_count = result if isinstance(result, int) else 0
            print(f"[ReportService] 已删除 {deleted_count} 条过期日报")
            return deleted_count
            
        except Exception as e:
            print(f"[ReportService] 删除过期日报失败: {e}")
            return 0
