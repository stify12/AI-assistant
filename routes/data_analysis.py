"""
数据分析路由模块
提供数据分析任务管理、工作流和报告生成功能
"""
import os
import io
import uuid
import json
import requests
from datetime import datetime
from flask import Blueprint, request, jsonify, render_template, Response

from services.config_service import ConfigService
from services.storage_service import StorageService
from services.llm_service import LLMService

data_analysis_bp = Blueprint('data_analysis', __name__)

ANALYSIS_TASKS_DIR = 'analysis_tasks'
ANALYSIS_FILES_DIR = 'analysis_files'


@data_analysis_bp.route('/data-analysis')
def data_analysis_page():
    """数据分析页面"""
    return render_template('data-analysis.html')


@data_analysis_bp.route('/api/analysis/tasks', methods=['GET', 'POST'])
def analysis_tasks_api():
    """分析任务列表和创建"""
    if request.method == 'GET':
        tasks = []
        for filename in StorageService.list_json_files(ANALYSIS_TASKS_DIR):
            task_id = filename[:-5]
            task = StorageService.load_analysis_task(task_id)
            if task:
                tasks.append({
                    'task_id': task_id,
                    'name': task.get('name', '未命名'),
                    'description': task.get('description', ''),
                    'status': task.get('status', 'pending'),
                    'file_count': len(task.get('files', [])),
                    'created_at': task.get('created_at'),
                    'updated_at': task.get('updated_at')
                })
        tasks.sort(key=lambda x: x.get('updated_at', ''), reverse=True)
        return jsonify(tasks)
    
    else:  # POST
        data = request.json
        task_id = str(uuid.uuid4())[:8]
        
        task_data = {
            'task_id': task_id,
            'name': data.get('name', '未命名任务'),
            'description': data.get('description', ''),
            'status': 'pending',
            'created_at': datetime.now().isoformat(),
            'updated_at': datetime.now().isoformat(),
            'files': [],
            'workflow_state': {
                'current_step': None,
                'steps': [
                    {'id': 'parse', 'name': '数据解析', 'status': 'pending'},
                    {'id': 'analyze', 'name': '内容分析', 'status': 'pending'},
                    {'id': 'template', 'name': '模板生成', 'status': 'pending'},
                    {'id': 'report', 'name': '报告编写', 'status': 'pending'}
                ]
            },
            'results': {}
        }
        
        StorageService.save_analysis_task(task_id, task_data)
        return jsonify({'task_id': task_id, 'created_at': task_data['created_at']})


@data_analysis_bp.route('/api/analysis/tasks/<task_id>', methods=['GET', 'DELETE'])
def analysis_task_detail(task_id):
    """任务详情和删除"""
    if request.method == 'GET':
        task = StorageService.load_analysis_task(task_id)
        if not task:
            return jsonify({'error': '任务不存在'}), 404
        return jsonify(task)
    
    else:  # DELETE
        task = StorageService.load_analysis_task(task_id)
        if task:
            # 删除关联文件
            for file_info in task.get('files', []):
                file_path = os.path.join(ANALYSIS_FILES_DIR, file_info.get('file_id', ''))
                if os.path.exists(file_path):
                    os.remove(file_path)
            # 删除任务
            StorageService.delete_file(StorageService.get_file_path(ANALYSIS_TASKS_DIR, task_id))
        return jsonify({'success': True})


@data_analysis_bp.route('/api/analysis/tasks/<task_id>/files', methods=['GET', 'POST'])
def analysis_files_api(task_id):
    """文件上传和列表"""
    task = StorageService.load_analysis_task(task_id)
    if not task:
        return jsonify({'error': '任务不存在'}), 404
    
    if request.method == 'GET':
        return jsonify(task.get('files', []))
    
    else:  # POST
        if 'files' not in request.files:
            return jsonify({'error': '未上传文件'}), 400
        
        files = request.files.getlist('files')
        uploaded = []
        errors = []
        
        StorageService.ensure_dir(ANALYSIS_FILES_DIR)
        
        for file in files:
            filename = file.filename
            
            if not (filename.lower().endswith('.xlsx') or filename.lower().endswith('.docx')):
                errors.append({'filename': filename, 'error': '仅支持 .xlsx 和 .docx 格式'})
                continue
            
            file.seek(0, 2)
            size = file.tell()
            file.seek(0)
            
            if size > 10 * 1024 * 1024:
                errors.append({'filename': filename, 'error': '文件大小超过10MB限制'})
                continue
            
            file_id = str(uuid.uuid4())[:8] + '_' + filename
            file_path = os.path.join(ANALYSIS_FILES_DIR, file_id)
            file.save(file_path)
            
            file_info = {
                'file_id': file_id,
                'filename': filename,
                'size': size,
                'upload_time': datetime.now().isoformat(),
                'status': 'uploaded',
                'parse_result': None,
                'analysis_result': None
            }
            
            task['files'].append(file_info)
            uploaded.append(file_info)
        
        StorageService.save_analysis_task(task_id, task)
        
        return jsonify({
            'uploaded': uploaded,
            'errors': errors,
            'total_files': len(task['files'])
        })


@data_analysis_bp.route('/api/analysis/tasks/<task_id>/workflow/start', methods=['POST'])
def start_analysis_workflow(task_id):
    """启动分析工作流"""
    task = StorageService.load_analysis_task(task_id)
    if not task:
        return jsonify({'error': '任务不存在'}), 404
    
    if not task.get('files'):
        return jsonify({'error': '请先上传文件'}), 400
    
    task['status'] = 'running'
    task['workflow_state']['current_step'] = 'parse'
    task['workflow_state']['steps'][0]['status'] = 'running'
    task['workflow_state']['steps'][0]['started_at'] = datetime.now().isoformat()
    
    StorageService.save_analysis_task(task_id, task)
    
    return jsonify({'status': 'started', 'current_step': 'parse'})


@data_analysis_bp.route('/api/analysis/tasks/<task_id>/workflow/status', methods=['GET'])
def get_workflow_status(task_id):
    """获取工作流状态"""
    task = StorageService.load_analysis_task(task_id)
    if not task:
        return jsonify({'error': '任务不存在'}), 404
    
    return jsonify({
        'status': task.get('status'),
        'workflow_state': task.get('workflow_state'),
        'files': [{
            'file_id': f.get('file_id'),
            'filename': f.get('filename'),
            'status': f.get('status'),
            'has_parse_result': f.get('parse_result') is not None,
            'has_analysis_result': f.get('analysis_result') is not None
        } for f in task.get('files', [])]
    })


@data_analysis_bp.route('/api/analysis/tasks/<task_id>/report', methods=['GET'])
def get_analysis_report(task_id):
    """获取分析报告"""
    task = StorageService.load_analysis_task(task_id)
    if not task:
        return jsonify({'error': '任务不存在'}), 404
    
    return jsonify({
        'template': task.get('results', {}).get('template'),
        'report': task.get('results', {}).get('report'),
        'files': [{
            'filename': f['filename'],
            'parse_result': f.get('parse_result'),
            'analysis_result': f.get('analysis_result')
        } for f in task.get('files', [])]
    })


@data_analysis_bp.route('/api/analysis/tasks/<task_id>/download', methods=['GET'])
def download_analysis_report(task_id):
    """下载Word报告"""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    task = StorageService.load_analysis_task(task_id)
    if not task:
        return jsonify({'error': '任务不存在'}), 404
    
    report = task.get('results', {}).get('report', {})
    if not report:
        return jsonify({'error': '报告尚未生成'}), 400
    
    content = report.get('content', '')
    
    doc = Document()
    
    title = doc.add_heading(task.get('name', '数据分析报告'), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"生成时间：{report.get('generated_at', '')}")
    doc.add_paragraph()
    
    for line in content.split('\n'):
        line = line.strip()
        if not line:
            continue
        
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    return Response(
        output.getvalue(),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment;filename=analysis_report_{task_id}.docx'}
    )
