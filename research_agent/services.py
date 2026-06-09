"""
AI调研员 - 服务层

包含：
- ModelService: LLM模型调用服务（复用项目已有模式）
- SearchService: 联网搜索服务（基于DuckDuckGo免费搜索）
- DataCleanService: 数据清洗与反幻觉校验
- FileParseService: 用户上传文件解析（CSV/Word/PDF/TXT）
- ReportService: 报告生成与导出（Markdown/Word）
- StorageService: 任务持久化存储
"""

import os
import re
import json
import uuid
import time
import requests
from typing import List, Dict, Any, Optional, Tuple, Generator
from datetime import datetime
from urllib.parse import quote_plus

from .models import (
    ResearchTask, TaskStatus, CollectedData, DataSource,
    DataSourceType, CredibilityLevel, ResearchReport, ReportFormat
)


# ============ LLM 模型服务 ============

class ModelService:
    """模型调用服务 - 封装对各LLM API的调用"""
    
    def __init__(self, config_path: str = "config.json", user_id: int = None):
        self.config_path = config_path
        self.user_id = user_id
        self._load_config()
    
    def _load_config(self):
        """加载配置 - 优先从数据库获取用户API密钥"""
        try:
            from services.config_service import ConfigService
            self.config = ConfigService.load_config(apply_headers=True, user_id=self.user_id)
        except Exception as e:
            print(f"[ModelService] ConfigService加载失败: {e}，回退到文件配置")
            try:
                with open(self.config_path, 'r', encoding='utf-8') as f:
                    self.config = json.load(f)
            except Exception:
                self.config = {}
    
    def _get_api_config(self, model_name: str) -> Tuple[str, str, str]:
        """
        根据模型名称获取对应的API地址、密钥和实际模型ID
        
        Returns:
            (api_url, api_key, actual_model_name)
        """
        if 'deepseek' in model_name.lower():
            return (
                "https://api.deepseek.com/v1/chat/completions",
                self.config.get("deepseek_api_key", ""),
                "deepseek-chat"  # DeepSeek API要求使用 deepseek-chat 作为模型名
            )
        elif 'gpt' in model_name.lower():
            return (
                self.config.get("gpt_api_url", "https://api.openai.com/v1/chat/completions"),
                self.config.get("gpt_api_key", ""),
                model_name
            )
        else:
            # 默认使用豆包/Ark API
            return (
                self.config.get("api_url", "https://ark.cn-beijing.volces.com/api/v3/chat/completions"),
                self.config.get("api_key", ""),
                model_name
            )
    
    def call_llm(self, model_name: str, system_prompt: str, user_prompt: str,
                 temperature: float = 0.3, max_tokens: int = 8192,
                 timeout: int = 120, max_retries: int = 2) -> Tuple[str, Dict]:
        """
        调用LLM模型（非流式）
        
        Args:
            model_name: 模型名称
            system_prompt: 系统提示词
            user_prompt: 用户提示词
            temperature: 温度参数
            max_tokens: 最大输出token数
            timeout: 超时时间
            max_retries: 最大重试次数
            
        Returns:
            (response_text, usage_info)
        """
        api_url, api_key, actual_model = self._get_api_config(model_name)
        
        if not api_key:
            raise ValueError(f"未配置模型 {model_name} 的API密钥，请在config.json中设置")
        
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}"
        }
        
        payload = {
            "model": actual_model,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            "temperature": temperature,
            "max_tokens": max_tokens
        }
        
        last_error = None
        for attempt in range(max_retries + 1):
            try:
                response = requests.post(
                    api_url, headers=headers, json=payload, timeout=timeout
                )
                response.raise_for_status()
                result = response.json()
                
                content = result.get("choices", [{}])[0].get("message", {}).get("content", "")
                usage = result.get("usage", {})
                
                return content, usage
                
            except requests.exceptions.Timeout:
                last_error = f"请求超时（第{attempt + 1}次）"
                if attempt < max_retries:
                    time.sleep(2 * (attempt + 1))
            except requests.exceptions.HTTPError as e:
                last_error = f"HTTP错误: {e}"
                if attempt < max_retries:
                    time.sleep(2 * (attempt + 1))
            except Exception as e:
                last_error = f"调用失败: {str(e)}"
                if attempt < max_retries:
                    time.sleep(1)
        
        raise Exception(f"LLM调用失败（已重试{max_retries}次）: {last_error}")
    
    def call_llm_stream(self, model_name: str, system_prompt: str, user_prompt: str,
                        temperature: float = 0.3, max_tokens: int = 8192,
                        timeout: int = 180) -> Generator:
        """
        流式调用LLM模型
        
        Yields:
            dict: {'type': 'content', 'content': str} 或 {'type': 'usage', 'usage': dict}
        """
        api_url, api_key, actual_model = self._get_api_config(model_name)
        
        if not api_key:
            raise ValueError(f"未配置模型 {model_name} 的API密钥")
        
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}"
        }
        
        payload = {
            "model": actual_model,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            "temperature": temperature,
            "max_tokens": max_tokens,
            "stream": True
        }
        
        response = requests.post(
            api_url, headers=headers, json=payload,
            timeout=timeout, stream=True
        )
        response.raise_for_status()
        
        for line in response.iter_lines():
            if not line:
                continue
            line_str = line.decode('utf-8')
            if line_str.startswith('data: '):
                data_str = line_str[6:]
                if data_str.strip() == '[DONE]':
                    break
                try:
                    chunk = json.loads(data_str)
                    delta = chunk.get("choices", [{}])[0].get("delta", {})
                    if "content" in delta and delta["content"]:
                        yield {"type": "content", "content": delta["content"]}
                    
                    # 某些API在最后一个chunk中返回usage
                    if "usage" in chunk and chunk["usage"]:
                        yield {"type": "usage", "usage": chunk["usage"]}
                except json.JSONDecodeError:
                    continue


# ============ 联网搜索服务 ============

class SearchService:
    """
    联网搜索服务 - 多搜索引擎支持
    
    搜索优先级：
    1. 百度搜索（中文搜索最佳，免费）
    2. DuckDuckGo（国际搜索，免费）
    3. 搜狗搜索（备选，免费）
    4. Bing API（付费，需要API Key）
    """
    
    # 多个User-Agent轮换使用
    USER_AGENTS = [
        'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
        'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:121.0) Gecko/20100101 Firefox/121.0',
        'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/17.2 Safari/605.1.15',
        'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36 Edg/120.0.0.0'
    ]
    
    # 搜索引擎URL
    DUCKDUCKGO_URL = "https://html.duckduckgo.com/html/"
    BAIDU_URL = "https://www.baidu.com/s"
    SOGOU_URL = "https://www.sogou.com/web"
    BING_API_URL = "https://api.bing.microsoft.com/v7.0/search"
    
    def __init__(self, config: dict = None):
        self.config = config or {}
        self._ua_index = 0
        self.session = requests.Session()
        self._update_headers()
    
    def _update_headers(self):
        """更新请求头，轮换User-Agent"""
        self._ua_index = (self._ua_index + 1) % len(self.USER_AGENTS)
        self.session.headers.update({
            'User-Agent': self.USER_AGENTS[self._ua_index],
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8',
            'Accept-Encoding': 'gzip, deflate',
            'Connection': 'keep-alive',
        })
    
    def search(self, query: str, max_results: int = 10) -> List[Dict[str, str]]:
        """
        执行联网搜索 - 多搜索引擎自动切换
        
        Args:
            query: 搜索关键词
            max_results: 最大结果数
            
        Returns:
            搜索结果列表 [{"title": "", "url": "", "snippet": "", "source": ""}]
        """
        all_results = []
        errors = []
        
        # 1. 优先百度搜索（中文内容最佳）
        try:
            self._update_headers()
            results = self._search_baidu(query, max_results)
            if results:
                print(f"[SearchService] 百度搜索成功，获取 {len(results)} 条结果")
                all_results.extend(results)
                if len(all_results) >= max_results:
                    return all_results[:max_results]
        except Exception as e:
            errors.append(f"百度: {e}")
            print(f"[SearchService] 百度搜索失败: {e}")
        
        # 2. DuckDuckGo搜索
        try:
            self._update_headers()
            results = self._search_duckduckgo(query, max_results)
            if results:
                print(f"[SearchService] DuckDuckGo搜索成功，获取 {len(results)} 条结果")
                all_results.extend(results)
                if len(all_results) >= max_results:
                    return self._dedupe_results(all_results)[:max_results]
        except Exception as e:
            errors.append(f"DuckDuckGo: {e}")
            print(f"[SearchService] DuckDuckGo搜索失败: {e}")
        
        # 3. 搜狗搜索（备选）
        try:
            self._update_headers()
            results = self._search_sogou(query, max_results)
            if results:
                print(f"[SearchService] 搜狗搜索成功，获取 {len(results)} 条结果")
                all_results.extend(results)
                if len(all_results) >= max_results:
                    return self._dedupe_results(all_results)[:max_results]
        except Exception as e:
            errors.append(f"搜狗: {e}")
            print(f"[SearchService] 搜狗搜索失败: {e}")
        
        # 4. Bing API（如果配置了Key）
        bing_key = self.config.get("bing_api_key", "")
        if bing_key:
            try:
                results = self._search_bing(query, max_results, bing_key)
                if results:
                    print(f"[SearchService] Bing搜索成功，获取 {len(results)} 条结果")
                    all_results.extend(results)
            except Exception as e:
                errors.append(f"Bing: {e}")
                print(f"[SearchService] Bing搜索失败: {e}")
        
        if not all_results and errors:
            print(f"[SearchService] 所有搜索引擎都失败了: {'; '.join(errors)}")
        
        return self._dedupe_results(all_results)[:max_results]
    
    def _dedupe_results(self, results: List[Dict]) -> List[Dict]:
        """去重搜索结果"""
        seen_urls = set()
        unique_results = []
        for r in results:
            url = r.get('url', '')
            # 简单URL标准化
            normalized_url = url.rstrip('/').lower()
            if normalized_url and normalized_url not in seen_urls:
                seen_urls.add(normalized_url)
                unique_results.append(r)
        return unique_results
    
    def _search_baidu(self, query: str, max_results: int) -> List[Dict[str, str]]:
        """百度搜索（免费，中文搜索最佳）"""
        results = []
        try:
            params = {
                'wd': query,
                'rn': min(max_results, 50),  # 每页最多50条
                'ie': 'utf-8'
            }
            
            resp = self.session.get(self.BAIDU_URL, params=params, timeout=15)
            resp.raise_for_status()
            resp.encoding = 'utf-8'
            html = resp.text
            
            # 百度搜索结果解析
            # 结果在 <div class="result c-container ..."> 或 <div class="c-container">
            import re
            
            # 方案1：提取标准搜索结果
            # 标题在 <h3 class="..."><a href="...">TITLE</a></h3>
            # 摘要在 <span class="content-right_...">SNIPPET</span> 或其他位置
            
            # 提取所有搜索结果块
            result_blocks = re.findall(
                r'<div[^>]*class="[^"]*result[^"]*c-container[^"]*"[^>]*>(.*?)</div>\s*(?=<div|$)',
                html, re.DOTALL
            )
            
            if not result_blocks:
                # 备用方案：直接提取链接
                result_blocks = re.findall(
                    r'<div[^>]*class="[^"]*c-container[^"]*"[^>]*>(.*?)</div>\s*(?=<div|$)',
                    html, re.DOTALL
                )
            
            for block in result_blocks[:max_results]:
                # 提取标题和URL
                title_match = re.search(
                    r'<h3[^>]*>.*?<a[^>]*href="([^"]*)"[^>]*>(.*?)</a>',
                    block, re.DOTALL
                )
                
                if title_match:
                    url = title_match.group(1)
                    title = re.sub(r'<[^>]+>', '', title_match.group(2)).strip()
                    
                    # 提取摘要
                    snippet = ""
                    snippet_match = re.search(
                        r'<span[^>]*class="[^"]*content[^"]*"[^>]*>(.*?)</span>',
                        block, re.DOTALL
                    )
                    if snippet_match:
                        snippet = re.sub(r'<[^>]+>', '', snippet_match.group(1)).strip()
                    else:
                        # 备用：提取任意长文本
                        text_match = re.search(r'>([^<]{50,})<', block)
                        if text_match:
                            snippet = text_match.group(1).strip()[:200]
                    
                    if title and url:
                        results.append({
                            "title": title,
                            "url": url,
                            "snippet": snippet,
                            "source": "baidu"
                        })
            
            # 如果上述方法失败，使用更宽松的匹配
            if not results:
                # 直接提取所有带href的h3标题
                h3_links = re.findall(
                    r'<h3[^>]*>\s*<a[^>]*href="([^"]*)"[^>]*>(.*?)</a>\s*</h3>',
                    html, re.DOTALL
                )
                for url, title in h3_links[:max_results]:
                    title = re.sub(r'<[^>]+>', '', title).strip()
                    if title and url and 'baidu.com' not in url:
                        results.append({
                            "title": title,
                            "url": url,
                            "snippet": "",
                            "source": "baidu"
                        })
            
        except Exception as e:
            raise Exception(f"百度搜索异常: {str(e)}")
        
        return results
    
    def _search_duckduckgo(self, query: str, max_results: int) -> List[Dict[str, str]]:
        """DuckDuckGo HTML搜索"""
        results = []
        try:
            # 添加随机延迟避免被封
            time.sleep(0.5)
            
            resp = self.session.post(
                self.DUCKDUCKGO_URL,
                data={"q": query, "b": "", "kl": "cn-zh"},  # 添加中文区域
                timeout=20
            )
            resp.raise_for_status()
            html = resp.text
            
            import re
            
            # 提取结果链接和标题
            link_pattern = r'<a[^>]*class="result__a"[^>]*href="([^"]*)"[^>]*>(.*?)</a>'
            snippet_pattern = r'<a[^>]*class="result__snippet"[^>]*>(.*?)</a>'
            
            links = re.findall(link_pattern, html, re.DOTALL)
            snippets = re.findall(snippet_pattern, html, re.DOTALL)
            
            for i, (url, title) in enumerate(links[:max_results]):
                title = re.sub(r'<[^>]+>', '', title).strip()
                snippet = ""
                if i < len(snippets):
                    snippet = re.sub(r'<[^>]+>', '', snippets[i]).strip()
                
                # DuckDuckGo的URL可能是重定向链接
                if url.startswith('//duckduckgo.com/l/'):
                    url_match = re.search(r'uddg=([^&]+)', url)
                    if url_match:
                        from urllib.parse import unquote
                        url = unquote(url_match.group(1))
                
                if title and url:
                    results.append({
                        "title": title,
                        "url": url,
                        "snippet": snippet,
                        "source": "duckduckgo"
                    })
            
        except Exception as e:
            raise Exception(f"DuckDuckGo搜索异常: {str(e)}")
        
        return results
    
    def _search_sogou(self, query: str, max_results: int) -> List[Dict[str, str]]:
        """搜狗搜索（免费备选）"""
        results = []
        try:
            params = {
                'query': query,
                'ie': 'utf8'
            }
            
            resp = self.session.get(self.SOGOU_URL, params=params, timeout=15)
            resp.raise_for_status()
            resp.encoding = 'utf-8'
            html = resp.text
            
            import re
            
            # 搜狗搜索结果格式
            # <h3 class="..."><a href="URL">TITLE</a></h3>
            h3_links = re.findall(
                r'<h3[^>]*>\s*<a[^>]*href="([^"]*)"[^>]*>(.*?)</a>',
                html, re.DOTALL
            )
            
            for url, title in h3_links[:max_results]:
                title = re.sub(r'<[^>]+>', '', title).strip()
                if title and url and 'sogou.com' not in url:
                    results.append({
                        "title": title,
                        "url": url,
                        "snippet": "",
                        "source": "sogou"
                    })
            
        except Exception as e:
            raise Exception(f"搜狗搜索异常: {str(e)}")
        
        return results
    
    def _search_bing(self, query: str, max_results: int, api_key: str) -> List[Dict[str, str]]:
        """Bing搜索API（付费备选）"""
        headers = {"Ocp-Apim-Subscription-Key": api_key}
        params = {"q": query, "count": max_results, "mkt": "zh-CN"}
        
        resp = requests.get(self.BING_API_URL, headers=headers, params=params, timeout=10)
        resp.raise_for_status()
        data = resp.json()
        
        results = []
        for item in data.get("webPages", {}).get("value", []):
            results.append({
                "title": item.get("name", ""),
                "url": item.get("url", ""),
                "snippet": item.get("snippet", ""),
                "source": "bing"
            })
        
        return results
    
    def fetch_page_content(self, url: str, max_length: int = 5000) -> str:
        """
        抓取网页正文内容（遵守robots协议）
        
        Args:
            url: 网页URL
            max_length: 最大提取字符数
            
        Returns:
            网页正文文本
        """
        try:
            # 先检查robots.txt
            if not self._check_robots(url):
                return "[该网页robots.txt禁止爬取]"
            
            resp = self.session.get(url, timeout=10)
            resp.raise_for_status()
            
            # 检测编码
            resp.encoding = resp.apparent_encoding or 'utf-8'
            html = resp.text
            
            # 移除script和style标签
            html = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL)
            html = re.sub(r'<style[^>]*>.*?</style>', '', html, flags=re.DOTALL)
            html = re.sub(r'<nav[^>]*>.*?</nav>', '', html, flags=re.DOTALL)
            html = re.sub(r'<footer[^>]*>.*?</footer>', '', html, flags=re.DOTALL)
            
            # 移除HTML标签，保留文本
            text = re.sub(r'<[^>]+>', ' ', html)
            # 清理空白
            text = re.sub(r'\s+', ' ', text).strip()
            
            return text[:max_length]
            
        except Exception as e:
            return f"[页面抓取失败: {str(e)}]"
    
    def _check_robots(self, url: str) -> bool:
        """简单检查robots.txt是否允许爬取"""
        try:
            from urllib.parse import urlparse
            parsed = urlparse(url)
            robots_url = f"{parsed.scheme}://{parsed.netloc}/robots.txt"
            
            resp = self.session.get(robots_url, timeout=5)
            if resp.status_code == 200:
                # 简单检查是否有Disallow: /
                content = resp.text.lower()
                # 如果全站禁止，则不爬取
                if 'disallow: /' in content and 'disallow: / ' not in content:
                    # 可能是全站禁止，但也可能只是某个路径
                    # 简单处理：如果有Allow规则，认为部分允许
                    if 'allow:' in content:
                        return True
                    return False
            return True
        except Exception:
            return True  # robots.txt不可访问时默认允许


# ============ 文件解析服务 ============

class FileParseService:
    """
    用户上传文件解析服务
    
    支持格式：CSV、Word(.docx)、PDF、TXT、JSON
    """
    
    ALLOWED_EXTENSIONS = {'csv', 'docx', 'doc', 'pdf', 'txt', 'json', 'xlsx', 'xls'}
    MAX_FILE_SIZE = 20 * 1024 * 1024  # 20MB
    
    @staticmethod
    def validate_file(filename: str, file_size: int) -> Tuple[bool, str]:
        """验证文件格式和大小"""
        if '.' not in filename:
            return False, "文件名缺少扩展名"
        
        ext = filename.rsplit('.', 1)[1].lower()
        if ext not in FileParseService.ALLOWED_EXTENSIONS:
            return False, f"不支持的文件格式: {ext}，支持 CSV/DOCX/PDF/TXT/JSON/XLSX"
        
        if file_size > FileParseService.MAX_FILE_SIZE:
            return False, f"文件大小超过限制 20MB"
        
        return True, ""
    
    @staticmethod
    def parse_file(filepath: str) -> str:
        """
        解析文件内容为文本
        
        Args:
            filepath: 文件路径
            
        Returns:
            解析后的文本内容
        """
        ext = filepath.rsplit('.', 1)[1].lower()
        
        try:
            if ext == 'txt':
                return FileParseService._parse_txt(filepath)
            elif ext == 'csv':
                return FileParseService._parse_csv(filepath)
            elif ext in ('docx', 'doc'):
                return FileParseService._parse_docx(filepath)
            elif ext == 'pdf':
                return FileParseService._parse_pdf(filepath)
            elif ext == 'json':
                return FileParseService._parse_json(filepath)
            elif ext in ('xlsx', 'xls'):
                return FileParseService._parse_excel(filepath)
            else:
                return f"[不支持的文件格式: {ext}]"
        except Exception as e:
            return f"[文件解析失败: {str(e)}]"
    
    @staticmethod
    def _parse_txt(filepath: str) -> str:
        """解析文本文件"""
        for encoding in ['utf-8', 'gbk', 'gb2312', 'utf-16']:
            try:
                with open(filepath, 'r', encoding=encoding) as f:
                    return f.read()
            except (UnicodeDecodeError, UnicodeError):
                continue
        return "[文件编码无法识别]"
    
    @staticmethod
    def _parse_csv(filepath: str) -> str:
        """解析CSV文件"""
        import csv
        rows = []
        for encoding in ['utf-8', 'gbk', 'gb2312']:
            try:
                with open(filepath, 'r', encoding=encoding) as f:
                    reader = csv.reader(f)
                    for i, row in enumerate(reader):
                        if i > 500:  # 最多读取500行
                            rows.append("... (更多数据已省略)")
                            break
                        rows.append(" | ".join(row))
                return "\n".join(rows)
            except (UnicodeDecodeError, UnicodeError):
                continue
        return "[CSV文件编码无法识别]"
    
    @staticmethod
    def _parse_docx(filepath: str) -> str:
        """解析Word文档"""
        try:
            from docx import Document
            doc = Document(filepath)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # 也提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    paragraphs.append(" | ".join(cells))
            
            return "\n\n".join(paragraphs)
        except ImportError:
            return "[需要安装python-docx库来解析Word文档]"
    
    @staticmethod
    def _parse_pdf(filepath: str) -> str:
        """解析PDF文件"""
        try:
            # 尝试使用PyPDF2
            from PyPDF2 import PdfReader
            reader = PdfReader(filepath)
            text_parts = []
            for page in reader.pages[:50]:  # 最多读取50页
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return "\n\n".join(text_parts)
        except ImportError:
            try:
                # 备选：使用pdfplumber
                import pdfplumber
                with pdfplumber.open(filepath) as pdf:
                    text_parts = []
                    for page in pdf.pages[:50]:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
                    return "\n\n".join(text_parts)
            except ImportError:
                return "[需要安装PyPDF2或pdfplumber库来解析PDF文件。pip install PyPDF2]"
    
    @staticmethod
    def _parse_json(filepath: str) -> str:
        """解析JSON文件"""
        with open(filepath, 'r', encoding='utf-8') as f:
            data = json.load(f)
        return json.dumps(data, ensure_ascii=False, indent=2)[:10000]
    
    @staticmethod
    def _parse_excel(filepath: str) -> str:
        """解析Excel文件"""
        try:
            from openpyxl import load_workbook
            wb = load_workbook(filepath, read_only=True)
            rows = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows.append(f"=== 工作表: {sheet_name} ===")
                for i, row in enumerate(ws.iter_rows(values_only=True)):
                    if i > 500:
                        rows.append("... (更多数据已省略)")
                        break
                    cells = [str(c) if c is not None else "" for c in row]
                    rows.append(" | ".join(cells))
            return "\n".join(rows)
        except Exception as e:
            return f"[Excel解析失败: {str(e)}]"


# ============ 数据清洗服务 ============

class DataCleanService:
    """
    数据清洗与反幻觉校验服务
    
    功能：
    - 去重去噪
    - 来源可信度评估
    - 数据冲突检测
    - 反幻觉校验
    """
    
    # 已知权威机构列表（用于可信度评估）
    AUTHORITATIVE_SOURCES = {
        'high': [
            '国家统计局', '中国人民银行', '工信部', '教育部', '科技部',
            '商务部', '财政部', '国务院', '发改委', '中国证监会',
            '世界银行', 'IMF', 'OECD', '联合国',
            '艾瑞咨询', 'IDC', 'Gartner', 'McKinsey', 'BCG',
            '中国互联网络信息中心', 'CNNIC', '中国信通院',
            '上市公司财报', '招股说明书', 'SEC'
        ],
        'medium': [
            '新华网', '人民网', '央视', '经济日报', '光明日报',
            '36氪', '虎嗅', '钛媒体', 'IT桔子',
            'Bloomberg', 'Reuters', 'Forbes', 'TechCrunch',
            '前瞻产业研究院', '观研天下', '弗若斯特沙利文',
            '知乎专栏', '微信公众号（认证机构）'
        ]
    }
    
    # 广告/营销关键词（用于过滤无效内容）
    SPAM_KEYWORDS = [
        '立即购买', '限时优惠', '点击领取', '免费试用', '加微信',
        '扫码关注', '粉丝福利', '转发抽奖', '薅羊毛',
        '100%', '保证赚钱', '日赚', '轻松月入'
    ]
    
    @staticmethod
    def assess_credibility(source_name: str, url: str = "") -> CredibilityLevel:
        """
        评估数据来源可信度
        
        Args:
            source_name: 来源名称
            url: 来源URL
            
        Returns:
            可信度等级
        """
        source_lower = source_name.lower()
        url_lower = url.lower() if url else ""
        
        # 检查是否为权威来源
        for src in DataCleanService.AUTHORITATIVE_SOURCES['high']:
            if src.lower() in source_lower or src.lower() in url_lower:
                return CredibilityLevel.HIGH
        
        for src in DataCleanService.AUTHORITATIVE_SOURCES['medium']:
            if src.lower() in source_lower or src.lower() in url_lower:
                return CredibilityLevel.MEDIUM
        
        # 检查URL域名
        if url:
            gov_domains = ['.gov.cn', '.gov', '.edu.cn', '.edu', '.org.cn', '.ac.cn']
            for domain in gov_domains:
                if domain in url_lower:
                    return CredibilityLevel.HIGH
        
        return CredibilityLevel.LOW
    
    @staticmethod
    def is_spam_content(content: str) -> bool:
        """检测是否为广告/营销内容"""
        for keyword in DataCleanService.SPAM_KEYWORDS:
            if keyword in content:
                return True
        return False
    
    @staticmethod
    def deduplicate(data_list: List[CollectedData], threshold: float = 0.8) -> List[CollectedData]:
        """
        数据去重
        
        使用简单的文本相似度比较去重
        """
        from difflib import SequenceMatcher
        
        unique_data = []
        seen_contents = []
        
        for item in data_list:
            is_dup = False
            for seen in seen_contents:
                ratio = SequenceMatcher(None, item.content[:200], seen[:200]).ratio()
                if ratio > threshold:
                    is_dup = True
                    break
            
            if not is_dup:
                unique_data.append(item)
                seen_contents.append(item.content)
        
        return unique_data
    
    @staticmethod
    def detect_conflicts(data_list: List[CollectedData]) -> List[Dict]:
        """
        检测数据冲突
        
        Returns:
            冲突列表 [{"data_ids": [...], "description": "..."}]
        """
        # 简单实现：基于标签分组，同一标签下的数据进行比较
        tag_groups = {}
        for item in data_list:
            for tag in item.tags:
                if tag not in tag_groups:
                    tag_groups[tag] = []
                tag_groups[tag].append(item)
        
        conflicts = []
        for tag, items in tag_groups.items():
            if len(items) > 1:
                # 标记为需要人工审核的潜在冲突
                conflicts.append({
                    "tag": tag,
                    "data_ids": [item.id for item in items],
                    "description": f"关于「{tag}」的数据有{len(items)}个来源，可能存在差异",
                    "items_summary": [
                        {"id": item.id, "source": item.source.name, "snippet": item.content[:100]}
                        for item in items
                    ]
                })
        
        return conflicts


# ============ 报告生成服务 ============

class ReportService:
    """
    报告生成与导出服务
    
    支持输出格式：Markdown、Word(.docx)、PPT大纲
    """
    
    @staticmethod
    def generate_word(report: ResearchReport, output_path: str) -> str:
        """
        将报告导出为Word文档
        
        Args:
            report: 调研报告对象
            output_path: 输出文件路径
            
        Returns:
            文件路径
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # 标题
            title_para = doc.add_heading(report.title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 生成日期
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = date_para.add_run(f"生成时间：{report.created_at} | 版本：v{report.version}")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(128, 128, 128)
            
            doc.add_paragraph()  # 空行
            
            # 执行摘要
            if report.summary:
                doc.add_heading("执行摘要", level=1)
                doc.add_paragraph(report.summary)
            
            # 各章节
            for section in report.sections:
                doc.add_heading(section.get("title", ""), level=1)
                content = section.get("content", "")
                
                # 简单Markdown转Word：处理标题和正文
                for line in content.split("\n"):
                    line = line.strip()
                    if not line:
                        continue
                    if line.startswith("### "):
                        doc.add_heading(line[4:], level=3)
                    elif line.startswith("## "):
                        doc.add_heading(line[3:], level=2)
                    elif line.startswith("# "):
                        doc.add_heading(line[2:], level=1)
                    elif line.startswith("- ") or line.startswith("* "):
                        doc.add_paragraph(line[2:], style='List Bullet')
                    elif re.match(r'^\d+\.', line):
                        doc.add_paragraph(line, style='List Number')
                    else:
                        # 处理加粗
                        para = doc.add_paragraph()
                        parts = re.split(r'(\*\*.*?\*\*)', line)
                        for part in parts:
                            if part.startswith('**') and part.endswith('**'):
                                run = para.add_run(part[2:-2])
                                run.bold = True
                            else:
                                para.add_run(part)
            
            # 附录
            if report.appendix:
                doc.add_page_break()
                doc.add_heading("附录", level=1)
                for appendix_item in report.appendix:
                    doc.add_heading(appendix_item.get("title", ""), level=2)
                    doc.add_paragraph(appendix_item.get("content", ""))
            
            # 保存
            os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
            doc.save(output_path)
            return output_path
            
        except ImportError:
            raise Exception("需要安装python-docx库：pip install python-docx")
    
    @staticmethod
    def generate_markdown(report: ResearchReport) -> str:
        """将报告导出为Markdown文本"""
        lines = []
        lines.append(f"# {report.title}")
        lines.append(f"\n> 生成时间：{report.created_at} | 版本：v{report.version}\n")
        
        if report.summary:
            lines.append("## 执行摘要\n")
            lines.append(report.summary)
            lines.append("")
        
        for section in report.sections:
            lines.append(f"\n## {section.get('title', '')}\n")
            lines.append(section.get("content", ""))
        
        if report.appendix:
            lines.append("\n---\n")
            lines.append("## 附录\n")
            for item in report.appendix:
                lines.append(f"### {item.get('title', '')}\n")
                lines.append(item.get("content", ""))
        
        return "\n".join(lines)
    
    @staticmethod
    def generate_ppt_outline(report: ResearchReport) -> str:
        """生成PPT大纲（Markdown格式）"""
        lines = []
        lines.append(f"# PPT大纲：{report.title}\n")
        lines.append("---\n")
        
        # 封面页
        lines.append("## 📄 封面页")
        lines.append(f"- 标题：{report.title}")
        lines.append(f"- 日期：{report.created_at}")
        lines.append("")
        
        # 目录页
        lines.append("## 📋 目录页")
        for i, section in enumerate(report.sections, 1):
            lines.append(f"- {i}. {section.get('title', '')}")
        lines.append("")
        
        # 执行摘要页
        if report.summary:
            lines.append("## 💡 执行摘要（1页）")
            lines.append(f"- {report.summary[:300]}")
            lines.append("")
        
        # 各章节页面建议
        for section in report.sections:
            title = section.get("title", "")
            content = section.get("content", "")
            
            lines.append(f"## 📊 {title}（建议1-3页）")
            
            # 提取关键要点
            key_points = []
            for line in content.split("\n"):
                line = line.strip()
                if line.startswith("- ") or line.startswith("* "):
                    key_points.append(line)
                elif line.startswith("**") and "**" in line[2:]:
                    key_points.append(f"- {line}")
            
            for point in key_points[:5]:
                lines.append(f"  {point}")
            
            lines.append("  - 💡 建议配图/图表类型：柱状图/饼图/表格")
            lines.append("")
        
        # 总结页
        lines.append("## 🎯 总结与建议（1页）")
        lines.append("- 核心结论要点")
        lines.append("- 行动建议")
        lines.append("")
        
        return "\n".join(lines)


# ============ 任务存储服务 ============

class StorageService:
    """任务持久化存储服务"""
    
    STORAGE_DIR = "research_tasks"
    
    def __init__(self):
        os.makedirs(self.STORAGE_DIR, exist_ok=True)
    
    def save_task(self, task: ResearchTask) -> str:
        """保存任务到JSON文件"""
        task.updated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        filepath = os.path.join(self.STORAGE_DIR, f"{task.task_id}.json")
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(task.to_dict(), f, ensure_ascii=False, indent=2)
        return filepath
    
    def load_task(self, task_id: str) -> Optional[ResearchTask]:
        """从JSON文件加载任务"""
        filepath = os.path.join(self.STORAGE_DIR, f"{task_id}.json")
        if not os.path.exists(filepath):
            return None
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                data = json.load(f)
            return ResearchTask.from_dict(data)
        except Exception as e:
            print(f"[StorageService] 加载任务失败: {e}")
            return None
    
    def list_tasks(self) -> List[Dict]:
        """列出所有任务摘要"""
        tasks = []
        for filename in os.listdir(self.STORAGE_DIR):
            if filename.endswith('.json'):
                filepath = os.path.join(self.STORAGE_DIR, filename)
                try:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                    tasks.append({
                        "task_id": data.get("task_id", ""),
                        "query": data.get("query", ""),
                        "status": data.get("status", ""),
                        "industry": data.get("industry", ""),
                        "created_at": data.get("created_at", ""),
                        "updated_at": data.get("updated_at", "")
                    })
                except Exception:
                    continue
        
        # 按更新时间倒序排列
        tasks.sort(key=lambda x: x.get("updated_at", ""), reverse=True)
        return tasks
    
    def delete_task(self, task_id: str) -> bool:
        """删除任务"""
        filepath = os.path.join(self.STORAGE_DIR, f"{task_id}.json")
        if os.path.exists(filepath):
            os.remove(filepath)
            return True
        return False
